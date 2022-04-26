# coding=utf-8
import re
from math import floor, ceil

from docx import Document

from 音标 import PhoneticSymbol

all_word_detail = []

file_name = r'D:/蒋一山的文件/英语/中考单词速记/单词/第33-34课.docx'
document = Document(file_name)
# 读取每段资料
# lines = [paragraph.text for paragraph in document.paragraphs]
# 输出并观察结果，也可以通过其他手段处理文本即可
# for one_line in lines:
for paragraph in document.paragraphs:
    one_line = paragraph.text
    one_line = one_line.strip()
    print(one_line)
    pattern = r'^([a-z A-Z]+)\s*(.*)'
    search_objs = re.search(pattern, one_line)
    if search_objs:
        w_english = search_objs.group(1)
        w_english = w_english.strip()
        w_chiese = search_objs.group(2)
        w_chiese = w_chiese.strip()
        print(w_english, w_chiese)

        ps = PhoneticSymbol()
        ps_w1 = ps.get_phonetic_symbol(w_english)

        new_one_line = f'{w_english} {ps_w1}'
        word_len = len(new_one_line)
        rejusted_word_len = ceil(word_len / 4) * 4
        tab_num = (28 - rejusted_word_len) / 4
        tab = f'\t' * int(tab_num)
        print(f'new_one_line:{new_one_line}, word_len:{word_len}, rejusted_word_len:{rejusted_word_len}, tab_num:{tab_num}, rejusted_word_len2:{rejusted_word_len+4*tab_num}')
        new_one_line = f'{w_english} {ps_w1}{tab}{w_chiese}'
        print(new_one_line)

        paragraph.text = new_one_line

document.save(r'D:/蒋一山的文件/英语/中考单词速记/单词/第33-34课单词_带音标.docx')


