# coding=utf-8
import re

from docx import Document

from 音标 import PhoneticSymbol

all_word_detail = []

file_name = r'音标练习.docx'
document = Document(file_name)
# 读取每段资料
# lines = [paragraph.text for paragraph in document.paragraphs]
# 输出并观察结果，也可以通过其他手段处理文本即可
# for one_line in lines:
for paragraph in document.paragraphs:
    one_line = paragraph.text
    print(one_line)
    # 匹配一：(    ) 1. pretty				spend				bread				desk
    # 匹配二：(    ) 1. A. look 			B. book 			C. over 			D. took
    pattern = r'^\(    \) [0-9]+\. (([A-D]\. )?([A-Za-z]+)[\s]*)(([A-D]\. )?([A-Za-z]+)[\s]*)(([A-D]\. )?([A-Za-z]+)[\s]*)(([A-D]\. )?([A-Za-z]+)[\s]*)'
    search_objs = re.search(pattern, one_line)
    if search_objs:
        w1 = search_objs.group(3)
        w2 = search_objs.group(6)
        w3 = search_objs.group(9)
        w4 = search_objs.group(12)
        print(w1, w2, w3, w4)

        ps = PhoneticSymbol()
        ps_w1 = ps.get_phonetic_symbol(w1)
        ps_w2 = ps.get_phonetic_symbol(w2)
        ps_w3 = ps.get_phonetic_symbol(w3)
        ps_w4 = ps.get_phonetic_symbol(w4)

        # 要匹配整个单词，不能用replace，否则可能会替换单词的一部分
        one_line = re.sub(f'{w1}\\b', f'{w1} {ps_w1}', one_line, 1)
        one_line = re.sub(f'{w2}\\b', f'{w2} {ps_w2}', one_line, 1)
        one_line = re.sub(f'{w3}\\b', f'{w3} {ps_w3}', one_line, 1)
        one_line = re.sub(f'{w4}\\b', f'{w4} {ps_w4}', one_line, 1)
        one_line = one_line.replace('\t\t\t', '\t')
        print(one_line)

        paragraph.text = one_line

document.save(r'音标练习_带音标.docx')


